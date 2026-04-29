"""
DOCX generator for Kreator CV — Anna Jakubowska.
Generates a clean, professional Word document from adapted CV data.
"""

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm


# ── Design constants ────────────────────────────────────────────────
COLOR_ACCENT  = RGBColor(0x1A, 0x3A, 0x5C)   # dark navy
COLOR_SECTION = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
COLOR_TEXT    = RGBColor(0x1F, 0x1F, 0x1F)   # near black
COLOR_LIGHT   = RGBColor(0x55, 0x55, 0x55)   # grey for dates

FONT_NAME = "Calibri"

LINKEDIN_URL = "https://www.linkedin.com/in/anna-jakubowska-market-researcher/"


def generate_cv_docx(cv_data: dict, output_path: Path) -> Path:
    """
    Generates a .docx CV from adapted cv_data (Anna Jakubowska).

    Args:
        cv_data:     Adapted CV dict (output of cv_adapter.adapt_cv()).
        output_path: Where to save the .docx file.

    Returns:
        output_path (for chaining).
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    _build_doc(cv_data).save(str(output_path))
    return output_path


def generate_cv_docx_bytes(cv_data: dict) -> bytes:
    """
    Generates a .docx CV entirely in memory and returns raw bytes.
    Avoids any filesystem dependency — safe for use in email sending.
    """
    buf = BytesIO()
    _build_doc(cv_data).save(buf)
    return buf.getvalue()


def _build_doc(cv_data: dict) -> Document:
    """Builds and returns a Document from adapted cv_data."""
    doc = Document()
    _set_margins(doc)

    # ── Header: Name ────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cv_data["personal"]["name"].upper())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_ACCENT
    run.font.name = FONT_NAME

    # ── Contact line: Phone | Email ──────────────────────────────────
    contact = f"{cv_data['personal']['phone']}  |  {cv_data['personal']['email']}"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(contact)
    r2.font.size = Pt(10)
    r2.font.color.rgb = COLOR_LIGHT
    r2.font.name = FONT_NAME

    # ── LinkedIn hyperlink ───────────────────────────────────────────
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_hyperlink(p3, "Mój LinkedIn", LINKEDIN_URL, Pt(10))

    # ── Job title (from posting) ─────────────────────────────────────
    job_title = cv_data.get("job_title", "").strip()
    if job_title:
        p_jt = doc.add_paragraph()
        p_jt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_jt.paragraph_format.space_before = Pt(12)
        p_jt.paragraph_format.space_after  = Pt(12)
        r_jt = p_jt.add_run(job_title.upper())
        r_jt.bold = True
        r_jt.font.size = Pt(18)
        r_jt.font.color.rgb = COLOR_ACCENT
        r_jt.font.name = FONT_NAME

    _add_divider(doc)

    lang = cv_data.get("cv_output_language", "pl")

    # ── Section headings (language-aware) ───────────────────────────
    HEADINGS = {
        "pl": {
            "summary":      "PODSUMOWANIE ZAWODOWE",
            "competencies": "KOMPETENCJE",
            "experience":   "DOŚWIADCZENIE ZAWODOWE",
            "education":    "WYKSZTAŁCENIE",
            "languages":    "ZNAJOMOŚĆ JĘZYKÓW",
            "interests":    "OBSZARY ZAINTERESOWAŃ",
        },
        "en-US": {
            "summary":      "PROFESSIONAL SUMMARY",
            "competencies": "COMPETENCIES",
            "experience":   "PROFESSIONAL EXPERIENCE",
            "education":    "EDUCATION",
            "languages":    "LANGUAGES",
            "interests":    "INTERESTS",
        },
    }
    H = HEADINGS.get(lang, HEADINGS["pl"])

    # ── Summary ─────────────────────────────────────────────────────
    _section_heading(doc, H["summary"])
    _body_paragraph(doc, cv_data.get("summary", ""))

    # ── Competencies ─────────────────────────────────────────────────
    _section_heading(doc, H["competencies"])
    comps = cv_data.get("competencies", [])
    if comps:
        comp_text = "  •  ".join(comps)
        _body_paragraph(doc, comp_text)

    # ── Work Experience ──────────────────────────────────────────────
    _section_heading(doc, H["experience"])
    fixed_facts = cv_data.get("fixed_experience_facts", [])
    if fixed_facts:
        bullets_lookup: dict[str, list] = {}
        for job in cv_data.get("experience", []):
            company_key = job.get("company", "").strip()
            if company_key:
                bullets_lookup[company_key] = job.get("bullets", [])

        for fact in fixed_facts:
            role    = fact.get("role_en") if lang == "en-US" else fact.get("role_pl", "")
            period  = fact.get("period_en") if lang == "en-US" else fact.get("period_pl", "")
            company  = fact["company"]
            industry = fact["industry"]
            bullets  = bullets_lookup.get(company, [])

            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = Pt(8)
            r_role = p_job.add_run(role)
            r_role.bold = True
            r_role.font.size = Pt(11)
            r_role.font.color.rgb = COLOR_SECTION
            r_role.font.name = FONT_NAME
            r_period = p_job.add_run(f"   {period}")
            r_period.font.size = Pt(10)
            r_period.font.color.rgb = COLOR_LIGHT
            r_period.font.name = FONT_NAME

            p_co = doc.add_paragraph()
            p_co.paragraph_format.space_before = Pt(0)
            p_co.paragraph_format.space_after  = Pt(2)
            r_co = p_co.add_run(f"{company}  |  {industry}")
            r_co.font.size = Pt(10)
            r_co.font.color.rgb = COLOR_LIGHT
            r_co.italic = True
            r_co.font.name = FONT_NAME

            for bullet in bullets:
                if bullet.strip():
                    p_b = doc.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.left_indent = Inches(0.25)
                    r_b = p_b.add_run(bullet)
                    r_b.font.size = Pt(10)
                    r_b.font.color.rgb = COLOR_TEXT
                    r_b.font.name = FONT_NAME
    else:
        for job in cv_data.get("experience", []):
            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = Pt(6)
            r_title = p_job.add_run(job["title"])
            r_title.bold = True
            r_title.font.size = Pt(11)
            r_title.font.color.rgb = COLOR_SECTION
            r_title.font.name = FONT_NAME
            r_dates = p_job.add_run(f"   {job['dates']}")
            r_dates.font.size = Pt(10)
            r_dates.font.color.rgb = COLOR_LIGHT
            r_dates.font.name = FONT_NAME

            for bullet in job.get("bullets", []):
                if bullet.strip():
                    p_b = doc.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.left_indent = Inches(0.25)
                    r_b = p_b.add_run(bullet)
                    r_b.font.size = Pt(10)
                    r_b.font.color.rgb = COLOR_TEXT
                    r_b.font.name = FONT_NAME

    # ── Education ───────────────────────────────────────────────────
    education = cv_data.get("education", [])
    if education:
        _section_heading(doc, H["education"])
        for edu in education:
            p_e = doc.add_paragraph()
            r_inst = p_e.add_run(edu["institution"])
            r_inst.bold = True
            r_inst.font.size = Pt(10)
            r_inst.font.name = FONT_NAME
            if edu.get("faculty"):
                p_e.add_run(f", {edu['faculty']}")
            if edu.get("degree"):
                p_e2 = doc.add_paragraph()
                p_e2.paragraph_format.left_indent = Inches(0.2)
                r_deg = p_e2.add_run(edu["degree"])
                r_deg.font.size = Pt(10)
                r_deg.font.color.rgb = COLOR_LIGHT
                r_deg.font.name = FONT_NAME

    # ── Languages ───────────────────────────────────────────────────
    _section_heading(doc, H["languages"])
    for lang_item in cv_data.get("languages", []):
        if lang == "en-US":
            name  = lang_item.get("language_en", lang_item["language"])
            level = lang_item.get("level_en", lang_item["level"])
        else:
            name  = lang_item["language"]
            level = lang_item["level"]
        _body_paragraph(doc, f"{name} – {level}")

    # ── Interests ───────────────────────────────────────────────────
    if cv_data.get("interests"):
        _section_heading(doc, H["interests"])
        _body_paragraph(doc, cv_data["interests"])

    # NOTE: ATS keywords are intentionally NOT rendered in the final CV.
    # They are used only as LLM context and shown in the UI analysis panel.

    # ── RODO clause ─────────────────────────────────────────────────
    if cv_data.get("rodo_clause"):
        p_rodo = doc.add_paragraph()
        p_rodo.paragraph_format.space_before = Pt(12)
        r_rodo = p_rodo.add_run(cv_data["rodo_clause"])
        r_rodo.font.size = Pt(7)
        r_rodo.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        r_rodo.font.name = FONT_NAME
        r_rodo.italic = True

    return doc


# ── Helpers ─────────────────────────────────────────────────────────

def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_ACCENT
    run.font.name = FONT_NAME
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A3A5C")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT
    run.font.name = FONT_NAME


def _add_hyperlink(paragraph, text: str, url: str, size: Pt | None = None) -> None:
    """Inserts a clickable hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(int(size.pt * 2)))
        rPr.append(sz)
        rPr.append(sz_cs)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)
